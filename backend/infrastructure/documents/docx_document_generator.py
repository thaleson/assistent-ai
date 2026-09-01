from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from backend.domain.enums.document_format import (
    DocumentFormat,
)
from backend.domain.interfaces.document_generator import (
    IDocumentGenerator,
)


class DocxDocumentGenerator(IDocumentGenerator):
    @property
    def document_format(self) -> DocumentFormat:
        """
        Return the DOCX document format.
        """
        return DocumentFormat.DOCX

    @property
    def media_type(self) -> str:
        """
        Return the DOCX MIME type.
        """
        return (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )

    @property
    def file_extension(self) -> str:
        """
        Return the DOCX file extension.
        """
        return "docx"

    def generate(
        self,
        title: str,
        content: str,
    ) -> bytes:
        """
        Generate a DOCX document in memory.
        """
        document = Document()

        normal_style = document.styles[
            "Normal"
        ]

        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        title_paragraph = (
            document.add_paragraph()
        )

        title_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        title_run = title_paragraph.add_run(
            title
        )

        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph()

        subtitle.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        subtitle_run = subtitle.add_run(
            "Resumo organizado pela Raissa AI",
        )

        subtitle_run.italic = True
        subtitle_run.font.size = Pt(9)

        document.add_paragraph()

        for raw_line in content.splitlines():
            line = raw_line.strip()

            if not line:
                document.add_paragraph()
                continue

            if line.startswith("### "):
                document.add_heading(
                    line[4:],
                    level=3,
                )
                continue

            if line.startswith("## "):
                document.add_heading(
                    line[3:],
                    level=2,
                )
                continue

            if line.startswith("# "):
                document.add_heading(
                    line[2:],
                    level=1,
                )
                continue

            if line.startswith(("- ", "* ")):
                document.add_paragraph(
                    line[2:],
                    style="List Bullet",
                )
                continue

            document.add_paragraph(line)

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer.getvalue()